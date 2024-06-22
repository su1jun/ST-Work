from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

from openpyxl import Workbook
from openpyxl.styles import Border, Side

from selenium import webdriver
import pyautogui
import pyperclip

import os
import numpy as np
import time
import cv2

import tkinter.messagebox as msgbox
from tkinter import * # __all__

class ChromeDriver:
    def __init__(self, loc, data, path):
        self.location = loc
        self.path = path
        self.data = data

        options = webdriver.ChromeOptions()
        self.driver = webdriver.Chrome(options=options)

    def errorImage(self, file_path, img_loc):    # 이미지 에러검증
        if img_loc:
            print(f"{file_path} 이미지가 발견되었습니다: {img_loc}")
            return True
        else:
            print(f"{file_path} 이미지를 찾을 수 없습니다.")
            return False
    
    def findImage(self, file_path, time=1):    # 이미지 위치 값 반환
        np_img = np.fromfile(os.path.join(self.path, 'data', 'images', file_path), np.uint8)
        img = cv2.imdecode(np_img, cv2.IMREAD_COLOR)
        img_loc = pyautogui.locateOnScreen(img, minSearchTime=time, confidence=0.8)
    
        if self.errorImage(file_path, img_loc):
            return img_loc
        else:
            return (0, 0)
        
    def findImageWait(self, file_path):    # 대기 걸고 이미지 위치 값 반환
        start_time = time.time()
        elapsed_time = time.time() - start_time
        timeout = 6
        np_img = np.fromfile(os.path.join(self.path, 'data', 'images', file_path), np.uint8)
        img = cv2.imdecode(np_img, cv2.IMREAD_COLOR)
        while True:
            img_loc = pyautogui.locateOnScreen(img, confidence=0.8)
            if img_loc:
                break

            elapsed_time = time.time() - start_time  # 경과 시간 계산
            if elapsed_time > timeout:
                print("시간 제한에 도달했습니다.")
                break

            time.sleep(0.1)

        print(f"로드바 이미지 감지 여부 : {elapsed_time}")

        if self.errorImage(file_path, img_loc):
            return img_loc
        else:
            return (0, 0)

    def clickImage(self, file_path, time = 2):    # 이미지 클릭
        self.location = self.findImage(file_path, time)
        pyautogui.click(self.location[0], self.location[1], button='left')
        return

    def drawInfo(self, wait_time):    # 페이지가 로드 되고 작업
        time.sleep(wait_time)
        try: # 팝업 창 예외 처리
            load_loc = self.findImage('load.png', wait_time) # load 글귀
            rec_off_loc = self.findImage('record_off.png') # record off

            if load_loc == (0, 0) or rec_off_loc == (0, 0):
                raise TypeError()

            pyautogui.click(rec_off_loc[0], rec_off_loc[1], button='left')
            pyautogui.moveTo(load_loc[0], load_loc[1])
            pyautogui.click(load_loc[0], load_loc[1], button='left')

        except TypeError:
            print("팝업 윈도우가 열렸습니다.")
            print(f"driver.window_handles : {self.driver.window_handles}")
            main_window_handle = self.driver.window_handles[0] # 현재 윈도우와 팝업 윈도우의 핸들을 가져옵니다.
            popup_window_handle = self.driver.window_handles[1]
            self.driver.switch_to.window(popup_window_handle) # 팝업 윈도우로 전환

            self.driver.close() # 팝업 윈도우 닫기

            self.driver.switch_to.window(main_window_handle) # 메인 웹페이지로 복귀
            pyautogui.press('f12') # 개발자 모드 다시 키고

            load_loc = self.findImage('load.png', wait_time) # load 글귀
            rec_off_loc = self.findImage('record_off.png') # record off

            if load_loc == (0, 0) or rec_off_loc == (0, 0) or load_loc == None or rec_off_loc == None:
                msgbox.showinfo('정보', '이미지 찾기를 실패하였습니다.')
            
            pyautogui.click(rec_off_loc[0], rec_off_loc[1], button='left')
            pyautogui.moveTo(load_loc[0], load_loc[1])
            pyautogui.click(load_loc[0], load_loc[1], button='left')

        finally:
            pyautogui.hotkey('ctrl', 'a')
            time.sleep(1)
            pyautogui.hotkey('ctrl', 'c')
            time.sleep(1)
            pyautogui.click(rec_off_loc[0], rec_off_loc[1], button='left')
            pyautogui.moveTo(10, 10)

            return pyperclip.paste()

    def run(self):
        # chrome driver setting
        # pyautogui.press('f12')  # 'f12' 키를 누름 -> 개발자 모드
        # self.clickImage('setting.png')
        # self.clickImage('popup.png')
        # self.clickImage('network.png')
        # self.clickImage('disable_cache.png')

        # site_path = os.path.join(self.path, 'data', 'setting', 'site_list.txt') # 저장된 사이트 목록 이름

        # # 사이트 목록 불러오고 저장하기
        # try:
        #     if os.path.exists(site_path):
        #         with open(site_path, 'r') as f:
        #             for line in f:
        #                 if ":" in line:
        #                     name, site = line.strip().split(":", 1)
        #                     target_sites[name.strip()] = site.strip()
        # except Exception as e:
        #     target_sites = {}
        #     print(f"parsing err : {e}")
            
        # if not target_sites:
        #     target_sites = {
        #         '서울과학기술대학교(국문)':'https://www.seoultech.ac.kr/index.jsp',
        #         '서울과학기술대학교(영문)':'https://en.seoultech.ac.kr/',
        #         '카이스트(국문)':'https://www.kaist.ac.kr/kr/',
        #         '카이스트(영문)':'https://www.kaist.ac.kr/en//',
        #         '서울시립대학교(국문)':'https://www.uos.ac.kr/main.do?epTicket=LOG',
        #         '서울시립대학교(영문)':'https://www.uos.ac.kr/en/main.do',
        #         '서울대학교(국문)':'https://www.snu.ac.kr/',
        #         '서울대학교(영문)':'https://en.snu.ac.kr/',
        #     }

        # try:
        #     with open(site_path, 'w' if os.path.exists(site_path) else 'w+') as f:
        #         for name, site in target_sites.items():
        #             f.write(f"{name}:{site}\n")
        #     print(f"file path: {os.path.abspath(site_path)}")
        # except Exception as e:
        #     print(f"save err : {e}")

        # # 웹사이트 성능 측정
        # performance_data = [[] for _ in range(len(target_sites))]

        test_times = int(self.data['test_times']) # 측정 횟수
        wait_time = int(self.data['wait_time']) # 로딩 대기 시간
        file_name = self.data['file_name'] # 저장 파일 이름
        save_path = self.data['save_path'] # 저장 파일 경로

        # for i in range(1, test_times + 1):
        #     lst_index = 0
        #     for site_name, site_url in target_sites.items():
        #         if i == 1:
        #             performance_data[lst_index].append(site_name)
        #         self.driver.get(site_url)
        #         try:
        #             temp = self.drawInfo(wait_time) # 성능 데이터
        #         except Exception as e:
        #             temp = ''
        #             print("성능 정보 불러오기 실패 : {e}")
        #         print(temp)

        #         performance_data[lst_index].append(temp)
        #         lst_index += 1
            
        #     print(performance_data)

        performance_data = [['서울과학기술대학교(국문)', '69 requests\r\n8.4 MB transferred\r\n8.9 MB resources\r\nFinish: 1.90 s\r\nDOMContentLoaded: 1.27 s\r\nLoad: 1.87 s', '64 requests\r\n8.4 MB transferred\r\n9.0 MB resources\r\nFinish: 1.41 s\r\nDOMContentLoaded: 1.04 s\r\nLoad: 1.45 s'], ['서울과학기술대학교(영문)', '51 requests\r\n7.0 MB transferred\r\n7.8 MB resources\r\nFinish: 937 ms\r\nDOMContentLoaded: 504 ms\r\nLoad: 888 ms', '48 requests\r\n7.0 MB transferred\r\n7.8 MB resources\r\nFinish: 638 ms\r\nDOMContentLoaded: 513 ms\r\nLoad: 668 ms'], ['카이스트(국문)', '44 requests\r\n3.6 MB transferred\r\n4.1 MB resources\r\nFinish: 3.46 s\r\nDOMContentLoaded: 1.63 s\r\nLoad: 2.92 s', '42 requests\r\n3.5 MB transferred\r\n3.9 MB resources\r\nFinish: 2.17 s\r\nDOMContentLoaded: 1.50 s\r\nLoad: 2.17 s'], ['카이스트(영문)', '37 requests\r\n4.4 MB transferred\r\n4.9 MB resources\r\nFinish: 2.25 s\r\nDOMContentLoaded: 843 ms\r\nLoad: 2.19 s', '35 requests\r\n4.4 MB transferred\r\n4.9 MB resources\r\nFinish: 2.30 s\r\nDOMContentLoaded: 1.57 s\r\nLoad: 1.57 s'], ['서울시립대학교(국문)', '99 requests\r\n12.7 MB transferred\r\n12.9 MB resources\r\nFinish: 3.12 s\r\nDOMContentLoaded: 574 ms\r\nLoad: 913 ms', '97 requests\r\n12.5 MB transferred\r\n12.6 MB resources\r\nFinish: 2.47 s\r\nDOMContentLoaded: 508 ms\r\nLoad: 773 ms'], ['서울시립대학교(영문)', '49 requests\r\n12.7 MB transferred\r\n13.0 MB resources\r\nFinish: 2.86 s\r\nDOMContentLoaded: 947 ms\r\nLoad: 1.36 s', '48 requests\r\n12.1 MB transferred\r\n12.4 MB resources\r\nFinish: 1.32 s\r\nDOMContentLoaded: 475 ms\r\nLoad: 827 ms'], ['서울대학교(국문)', '9 requests\r\n384 kB transferred\r\n379 kB resources\r\nFinish: 193 ms\r\nDOMContentLoaded: 57 ms\r\nLoad: 178 ms', '9 requests\r\n384 kB transferred\r\n379 kB resources\r\nFinish: 223 ms\r\nDOMContentLoaded: 157 ms\r\nLoad: 188 ms'], ['서울대학교(영문)', '56 requests\r\n11.7 MB transferred\r\n11.8 MB resources\r\nFinish: 3.33 s\r\nDOMContentLoaded: 528 ms\r\nLoad: 1.05 s', '53 requests\r\n8.3 MB transferred\r\n8.4 MB resources\r\nFinish: 3.35 s\r\nDOMContentLoaded: 569 ms\r\nLoad: 1.10 s']]

        file_name = os.path.join(save_path, file_name)
        self.savePerformanceData(file_name, performance_data)

        # self.driver.quit() # 웹드라이버 종료
        # return performance_data

        return [['서울과학기술대학교(국문)', '69 requests\r\n8.4 MB transferred\r\n8.9 MB resources\r\nFinish: 1.90 s\r\nDOMContentLoaded: 1.27 s\r\nLoad: 1.87 s', '64 requests\r\n8.4 MB transferred\r\n9.0 MB resources\r\nFinish: 1.41 s\r\nDOMContentLoaded: 1.04 s\r\nLoad: 1.45 s'], ['서울과학기술대학교(영문)', '51 requests\r\n7.0 MB transferred\r\n7.8 MB resources\r\nFinish: 937 ms\r\nDOMContentLoaded: 504 ms\r\nLoad: 888 ms', '48 requests\r\n7.0 MB transferred\r\n7.8 MB resources\r\nFinish: 638 ms\r\nDOMContentLoaded: 513 ms\r\nLoad: 668 ms'], ['카이스트(국문)', '44 requests\r\n3.6 MB transferred\r\n4.1 MB resources\r\nFinish: 3.46 s\r\nDOMContentLoaded: 1.63 s\r\nLoad: 2.92 s', '42 requests\r\n3.5 MB transferred\r\n3.9 MB resources\r\nFinish: 2.17 s\r\nDOMContentLoaded: 1.50 s\r\nLoad: 2.17 s'], ['카이스트(영문)', '37 requests\r\n4.4 MB transferred\r\n4.9 MB resources\r\nFinish: 2.25 s\r\nDOMContentLoaded: 843 ms\r\nLoad: 2.19 s', '35 requests\r\n4.4 MB transferred\r\n4.9 MB resources\r\nFinish: 2.30 s\r\nDOMContentLoaded: 1.57 s\r\nLoad: 1.57 s'], ['서울시립대학교(국문)', '99 requests\r\n12.7 MB transferred\r\n12.9 MB resources\r\nFinish: 3.12 s\r\nDOMContentLoaded: 574 ms\r\nLoad: 913 ms', '97 requests\r\n12.5 MB transferred\r\n12.6 MB resources\r\nFinish: 2.47 s\r\nDOMContentLoaded: 508 ms\r\nLoad: 773 ms'], ['서울시립대학교(영문)', '49 requests\r\n12.7 MB transferred\r\n13.0 MB resources\r\nFinish: 2.86 s\r\nDOMContentLoaded: 947 ms\r\nLoad: 1.36 s', '48 requests\r\n12.1 MB transferred\r\n12.4 MB resources\r\nFinish: 1.32 s\r\nDOMContentLoaded: 475 ms\r\nLoad: 827 ms'], ['서울대학교(국문)', '9 requests\r\n384 kB transferred\r\n379 kB resources\r\nFinish: 193 ms\r\nDOMContentLoaded: 57 ms\r\nLoad: 178 ms', '9 requests\r\n384 kB transferred\r\n379 kB resources\r\nFinish: 223 ms\r\nDOMContentLoaded: 157 ms\r\nLoad: 188 ms'], ['서울대학교(영문)', '56 requests\r\n11.7 MB transferred\r\n11.8 MB resources\r\nFinish: 3.33 s\r\nDOMContentLoaded: 528 ms\r\nLoad: 1.05 s', '53 requests\r\n8.3 MB transferred\r\n8.4 MB resources\r\nFinish: 3.35 s\r\nDOMContentLoaded: 569 ms\r\nLoad: 1.10 s']]
    
    def savePerformanceData(self, save_path, result):

    ##
    #     status_times = len(result[0]) - 1
    #     doc = Document()

    #     sections = doc.sections
    #     for section in sections:
    #         section.top_margin = Inches(0.5)  # 상단 여백
    #         section.bottom_margin = Inches(0.5)  # 하단 여백
    #         section.left_margin = Inches(0.5)  # 왼쪽 여백
    #         section.right_margin = Inches(0.5)  # 오른쪽 여백

    #     # doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    #     style = doc.styles['Normal']
    #     font = style.font
    #     font.name = '맑은 고딕'
    #     font.size = Pt(10)

    #     # 줄 간격 설정
    #     paragraph_format = style.paragraph_format
    #     paragraph_format.line_spacing = 0.8  # 단일 줄 간격
        
    #     for item in result:
    #         for i, texts in enumerate(item):
    #             # 대학 이름을 제목으로 추가
    #             if i == 0:
    #                 doc.add_paragraph(texts)
    #                 table = doc.add_table(rows=status_times, cols=2)

    #                 for row in table.rows:
    #                     for cell in row.cells:
    #                         self.set_cell_border(cell, start='single', end='single', top='single', bottom='single')
    #             else:
    #                 texts = texts.replace("\r\n", "\n")
    #                 table.cell(i - 1, 1).text = texts

    #         doc.add_paragraph("\n")

    #     doc.save(f"{save_path}.docx")

    #     ######
    #     # 엑셀 파일을 생성하고 활성 시트를 가져옵니다.
    #     wb = Workbook()
    #     ws = wb.active

    #     # 엑셀에서 사용할 테두리 스타일을 정의합니다.
    #     thin_border_side = Side(style='thin')
    #     thin_border = Border(
    #         left=thin_border_side,
    #         right=thin_border_side,
    #         top=thin_border_side,
    #         bottom=thin_border_side
    #     )

    #     start_row = 1
    #     start_col = 1
    #     for item in result:
    #         ws.cell(row=start_row, column=start_col, value=item[0])  # 대학 이름
    #         start_row += 2  # 성능 데이터를 위해 한 행을 건너뜁니다.

    #         # 성능 데이터를 입력합니다.
    #         for i, stat in enumerate(item[1:], start=start_row):
    #             ws.cell(row=i, column=start_col, value=stat)
    #             # 테두리 설정: 처음과 마지막 셀에만 테두리를 적용합니다.
    #             if i == start_row:
    #                 ws.cell(row=i, column=start_col).border = Border(top=thin_border_side, left=thin_border_side, right=thin_border_side)
    #             elif i == start_row + len(item[1:]) - 1:
    #                 ws.cell(row=i, column=start_col).border = Border(bottom=thin_border_side, left=thin_border_side, right=thin_border_side)
    #             else:
    #                 ws.cell(row=i, column=start_col).border = Border(left=thin_border_side, right=thin_border_side)
            
    #         start_row += len(item[1:]) + 1  # 다음 데이터 블록을 위한 줄바꿈


    #     # 데이터에 대한 테두리 설정
    #     for row in ws.iter_rows(min_row=start_row+2, min_col=start_col,
    #                             max_row=start_row+11, max_col=start_col):
    #         for cell in row:
    #             cell.border = thin_border

    #     # 열 너비 조절
    #     ws.column_dimensions['A'].width = 120
    #     ws.column_dimensions['B'].width = 120

    #     # 파일을 저장합니다.
    #     wb.save(f"{save_path}.xlsx")

    # def set_cell_border(self, cell, **kwargs):
    #     tc = cell._tc
    #     tcPr = tc.get_or_add_tcPr()

    #     for key, value in kwargs.items():
    #         tag = key.split('_')[0]  # start, end, top, bottom
    #         side = OxmlElement(f'w:{tag}')
    #         side.set(qn('w:val'), value)
    #         side.set(qn('w:sz'), '4')  # 테두리의 크기
    #         side.set(qn('w:color'), 'auto')
    #         side.set(qn('w:space'), '0')
    #         tcPr.append(side)